from docx import Document


class ResumeWriter:

    def __init__(self, resume_path):
        self.resume_path = resume_path
        self.document = Document(resume_path)

    def update_summary(self, new_summary):

        found_summary = False

        for paragraph in self.document.paragraphs:

            if paragraph.text.strip().upper() == "PROFESSIONAL SUMMARY":
                found_summary = True
                continue

            if found_summary:
                if paragraph.text.strip() != "":
                    paragraph.text = new_summary
                    break

    def update_skills(self, new_skills):

        found_skills = False

        for paragraph in self.document.paragraphs:

            if paragraph.text.strip().upper() == "TECHNICAL SKILLS":
                found_skills = True
                continue

            if found_skills:

                if paragraph.text.strip().upper() == "PROFESSIONAL EXPERIENCE":
                    break

                if paragraph.text.strip() == "":
                    continue

                paragraph.text = ""

        skills_text = ", ".join(new_skills)

        for i, paragraph in enumerate(self.document.paragraphs):

            if paragraph.text.strip().upper() == "TECHNICAL SKILLS":

                insert_para = self.document.paragraphs[i + 1]

                insert_para.text = skills_text
                break

    def save(self, output_path):
        self.document.save(output_path)